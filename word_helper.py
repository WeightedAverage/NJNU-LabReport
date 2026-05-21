#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Word报告生成辅助工具
处理序号重置和分节格式保持
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
import copy

class WordReportHelper:
    def __init__(self, template_path):
        self.template_path = template_path
        self.doc = Document(template_path)

    def reset_section_numbering(self, section_index, start_number=1):
        """
        重置指定节的序号
        section_index: 节索引（从0开始）
        start_number: 起始序号
        """
        # 遍历段落，找到属于该节的段落并重置序号
        # 注意：python-docx对节的支持有限，需要手动处理
        pass

    def add_section_with_same_format(self, section_content, preserve_header_footer=True):
        """
        添加新节，保持相同格式
        """
        # 获取当前最后一节的格式
        last_section = self.doc.sections[-1]

        # 添加新节
        new_section = self.doc.add_section()

        # 复制页边距
        new_section.top_margin = last_section.top_margin
        new_section.bottom_margin = last_section.bottom_margin
        new_section.left_margin = last_section.left_margin
        new_section.right_margin = last_section.right_margin

        # 复制页眉页脚
        if preserve_header_footer:
            # 复制页眉
            for para in last_section.header.paragraphs:
                new_para = new_section.header.add_paragraph(para.text)
                new_para.style = para.style

            # 复制页脚
            for para in last_section.footer.paragraphs:
                new_para = new_section.footer.add_paragraph(para.text)
                new_para.style = para.style

        return new_section

    def add_numbered_paragraph(self, text, level=1, start_new_number=False, number=1):
        """
        添加带序号的段落
        level: 序号级别（1=一、二、三，2=1.2.3，3=(1)(2)(3)）
        start_new_number: 是否重新从1开始
        number: 起始序号
        """
        # 根据级别生成序号
        if level == 1:
            chinese_nums = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
            prefix = f'{chinese_nums[number-1]}、' if number <= 10 else f'{number}.'
        elif level == 2:
            prefix = f'{number}.'
        else:
            prefix = f'({number})'

        # 添加段落
        para = self.doc.add_paragraph()
        run = para.add_run(f'{prefix}{text}')

        return para

    def save(self, output_path):
        """保存文档"""
        self.doc.save(output_path)
        print(f"已保存到: {output_path}")

def create_template():
    """创建带分节的Word模板"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)

    # 添加第一个节的内容
    para = doc.add_paragraph()
    run = para.add_run('一、实验目的')
    run.bold = True

    para = doc.add_paragraph()
    para.add_run('在此填写实验目的...')

    # 添加分页符，保持同一节格式
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(WD_BREAK.PAGE)

    # 添加第二个节
    para = doc.add_paragraph()
    run = para.add_run('二、实验原理')
    run.bold = True

    para = doc.add_paragraph()
    para.add_run('在此填写实验原理...')

    # 保存模板
    doc.save('01-原始模板/实验报告模板_分节版.docx')
    print("已创建分节模板")

if __name__ == "__main__":
    print("Word报告生成工具")
    print("1. 创建分节模板")
    print("2. 测试模板")

    choice = input("请选择: ").strip()

    if choice == '1':
        create_template()
    else:
        print("请运行 python generate_report.py")
