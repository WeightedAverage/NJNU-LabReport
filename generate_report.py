#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
电子专业实验报告生成脚本
适用于数字电子技术、模拟电子技术、单片机、嵌入式等课程
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')

from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

def read_experiment_info():
    """读取实验信息"""
    info_file = Path("实验信息.txt")
    if info_file.exists():
        return info_file.read_text(encoding='utf-8')
    return None

def parse_experiment_info(text):
    """解析实验信息"""
    data = {
        'name': '',
        'purpose': '',
        'principle': '',
        'hardware': '',
        'software': '',
        'results': '',
        'summary': ''
    }

    lines = text.split('\n')
    current_section = None

    for line in lines:
        if line.startswith('## 实验') and '：' in line:
            data['name'] = line.split('：')[-1].strip()
        elif line.startswith('## 二、实验目的') or line.startswith('## 实验目的'):
            current_section = 'purpose'
        elif line.startswith('## 三、实验原理') or line.startswith('## 实验原理'):
            current_section = 'principle'
        elif line.startswith('## 四、硬件电路设计') or line.startswith('## 硬件电路设计'):
            current_section = 'hardware'
        elif line.startswith('## 五、软件设计与代码实现') or line.startswith('## 软件设计'):
            current_section = 'software'
        elif line.startswith('## 六、实验结果与分析') or line.startswith('## 实验结果'):
            current_section = 'results'
        elif line.startswith('## 七、实验总结与心得') or line.startswith('## 实验总结'):
            current_section = 'summary'
        elif line.startswith('## ') or line.startswith('# '):
            current_section = None
        elif current_section and line.strip():
            data[current_section] += line + '\n'

    return data

def add_paragraph_after(doc, ref_paragraph, text, bold=False):
    """在指定段落后插入新段落"""
    new_para = doc.add_paragraph()
    new_para.text = text
    if bold:
        for run in new_para.runs:
            run.bold = True
    ref_paragraph._element.addnext(new_para._element)
    return new_para

def add_table_after(doc, ref_paragraph, headers, rows):
    """在指定段落后插入真实Word表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = cell_text

    ref_paragraph._element.addnext(table._element)
    return table

def add_code_block(doc, ref_paragraph, code_text):
    """添加代码块（等宽字体）"""
    code_para = doc.add_paragraph()
    code_para.text = code_text
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    ref_paragraph._element.addnext(code_para._element)
    return code_para

def generate_word_report(experiment_data):
    """生成Word报告"""
    template_dir = Path("01-模板")
    template_files = list(template_dir.glob("*.docx"))

    if not template_files:
        print("错误：01-模板/ 文件夹中没有找到.docx模板文件")
        return False

    template_path = template_files[0]
    print(f"使用模板：{template_path.name}")

    doc = Document(str(template_path))

    # 查找关键标记段落
    purpose_marker = None
    process_marker = None
    conclusion_marker = None

    for para in doc.paragraphs:
        text = para.text
        if '【实验预习】' in text:
            purpose_marker = para
        elif '【实验过程】' in text:
            process_marker = para
        elif '【结论与分析】' in text:
            conclusion_marker = para

    # 填充实验名称
    for para in doc.paragraphs[:10]:
        for run in para.runs:
            if '实验名称' in run.text:
                run.text = run.text.replace('实验名称', f'实验名称  {experiment_data["name"]}')

    # 【实验预习】部分
    if purpose_marker:
        last_para = purpose_marker

        # 实验目的
        purpose_title = add_paragraph_after(doc, last_para, '一、实验目的', bold=True)
        last_para = purpose_title

        purpose_content = experiment_data['purpose'].strip()
        if purpose_content:
            for line in purpose_content.split('\n'):
                line = line.strip()
                if line and not line.startswith('##'):
                    content_para = add_paragraph_after(doc, last_para, line)
                    last_para = content_para

        # 实验原理
        principle_title = add_paragraph_after(doc, last_para, '二、实验原理', bold=True)
        last_para = principle_title

        principle_content = experiment_data['principle'].strip()
        if principle_content:
            for line in principle_content.split('\n'):
                line = line.strip()
                if line and not line.startswith('##'):
                    content_para = add_paragraph_after(doc, last_para, line)
                    last_para = content_para

    # 【实验过程】部分
    if process_marker:
        last_para = process_marker

        # 实验器材
        equipment_title = add_paragraph_after(doc, last_para, '一、实验器材', bold=True)
        last_para = equipment_title

        equipment_items = [
            '• 开发板/实验箱',
            '• 下载器/仿真器',
            '• 数据线',
            '• 杜邦线若干',
            '• 电脑（安装开发软件）'
        ]
        for item in equipment_items:
            content_para = add_paragraph_after(doc, last_para, item)
            last_para = content_para

        # 硬件电路设计
        hw_title = add_paragraph_after(doc, last_para, '二、硬件电路设计', bold=True)
        last_para = hw_title

        hw_content = experiment_data['hardware'].strip()
        if hw_content:
            for line in hw_content.split('\n'):
                line = line.strip()
                if line and not line.startswith('##'):
                    content_para = add_paragraph_after(doc, last_para, line)
                    last_para = content_para

        # 图片占位
        img_title1 = add_paragraph_after(doc, last_para, '【此处插入硬件电路连接图】')
        last_para = img_title1

        # 实验步骤
        steps_title = add_paragraph_after(doc, last_para, '三、实验步骤', bold=True)
        last_para = steps_title

        steps = [
            '1. 打开开发环境，导入实验工程',
            '2. 编写驱动代码',
            '3. 编写主程序，实现功能',
            '4. 编译工程，确保无错误无警告',
            '5. 下载程序到开发板',
            '6. 观察实验现象，记录结果'
        ]
        for step in steps:
            content_para = add_paragraph_after(doc, last_para, step)
            last_para = content_para

        # 图片占位
        img_title2 = add_paragraph_after(doc, last_para, '【此处插入编译结果截图】')
        last_para = img_title2

        # 软件设计与代码实现
        sw_title = add_paragraph_after(doc, last_para, '四、软件设计与代码实现', bold=True)
        last_para = sw_title

        sw_content = experiment_data['software'].strip()
        if sw_content:
            # 检测是否包含代码块
            if '```c' in sw_content or '```python' in sw_content or '```:' in sw_content:
                # 提取代码块内容
                in_code = False
                code_lines = []
                for line in sw_content.split('\n'):
                    if line.strip().startswith('```'):
                        if in_code:
                            # 代码块结束，添加代码
                            code_text = '\n'.join(code_lines)
                            if code_text.strip():
                                last_para = add_code_block(doc, last_para, code_text)
                            code_lines = []
                            in_code = False
                        else:
                            in_code = True
                    elif in_code:
                        code_lines.append(line)
                    else:
                        line = line.strip()
                        if line and not line.startswith('##'):
                            content_para = add_paragraph_after(doc, last_para, line)
                            last_para = content_para
            else:
                # 普通文本，按行添加
                for line in sw_content.split('\n'):
                    line = line.strip()
                    if line and not line.startswith('##'):
                        content_para = add_paragraph_after(doc, last_para, line)
                        last_para = content_para

    # 【结论与分析】部分
    if conclusion_marker:
        last_para = conclusion_marker

        # 实验结果
        result_title = add_paragraph_after(doc, last_para, '一、实验结果与分析', bold=True)
        last_para = result_title

        result_content = experiment_data['results'].strip()
        if result_content:
            for line in result_content.split('\n'):
                line = line.strip()
                if line and not line.startswith('##'):
                    content_para = add_paragraph_after(doc, last_para, line)
                    last_para = content_para

        # 图片占位
        img_title3 = add_paragraph_after(doc, last_para, '【此处插入实验现象截图】')
        last_para = img_title3

        # 实验总结
        summary_title = add_paragraph_after(doc, last_para, '二、实验总结与心得', bold=True)
        last_para = summary_title

        summary_content = experiment_data['summary'].strip()
        if summary_content:
            for line in summary_content.split('\n'):
                line = line.strip()
                if line and not line.startswith('##'):
                    content_para = add_paragraph_after(doc, last_para, line)
                    last_para = content_para

    # 保存
    output_dir = Path("02-生成的报告")
    output_dir.mkdir(exist_ok=True)

    import random
    temp_filename = f"temp_{random.randint(1000, 9999)}.docx"
    temp_path = output_dir / temp_filename
    final_path = output_dir / f"{experiment_data['name']}.docx"

    try:
        doc.save(str(temp_path))

        if final_path.exists():
            try:
                final_path.unlink()
            except PermissionError:
                print("警告：旧文件被占用，将使用新文件名")
                final_path = output_dir / f"{experiment_data['name']}_v2.docx"

        temp_path.rename(final_path)
        print(f"Word报告已保存到: {final_path}")
        return True
    except Exception as e:
        print(f"错误：{e}")
        return False

def generate_markdown_outline(experiment_data):
    """生成Markdown大纲"""
    content = f"""# {experiment_data['name']}

## 实验基本信息
- 实验名称：{experiment_data['name']}

---

## 一、实验目的

{experiment_data['purpose'].strip() if experiment_data['purpose'].strip() else '（请填写实验目的）'}

---

## 二、实验原理

{experiment_data['principle'].strip() if experiment_data['principle'].strip() else '（请填写实验原理）'}

---

## 三、硬件电路设计

{experiment_data['hardware'].strip() if experiment_data['hardware'].strip() else '（请填写硬件电路设计）'}

【此处插入硬件电路连接图】

---

## 四、软件设计与代码实现

{experiment_data['software'].strip() if experiment_data['software'].strip() else '（请填写软件设计与代码）'}

---

## 五、实验结果与分析

{experiment_data['results'].strip() if experiment_data['results'].strip() else '（请填写实验结果）'}

【此处插入实验现象截图】

---

## 六、实验总结与心得

{experiment_data['summary'].strip() if experiment_data['summary'].strip() else '（请填写实验总结）'}

---

## 七、图片素材

- 电路图：请从 03-实验素材/电路图/ 选取
- 实验截图：请从 03-实验素材/实验截图/ 选取
- 数据波形：请从 03-实验素材/数据波形/ 选取
"""

    output_dir = Path("04-大纲草稿")
    output_dir.mkdir(exist_ok=True)
    output_path = output_dir / f"{experiment_data['name']}-大纲.md"

    output_path.write_text(content, encoding='utf-8')
    print(f"Markdown大纲已保存到: {output_path}")
    return True

def main():
    print("=" * 50)
    print("电子专业实验报告生成工具")
    print("=" * 50)

    info = read_experiment_info()
    if not info:
        print("错误：找不到实验信息.txt")
        print("请创建实验信息.txt或把讲义放到05-实验讲义/文件夹")
        return

    data = parse_experiment_info(info)
    if not data['name']:
        print("错误：无法解析实验名称")
        print("请确保实验信息.txt中包含 '## 实验X：实验名称' 格式")
        return

    print(f"\n实验名称：{data['name']}")

    print("\n请选择生成方式：")
    print("1. 生成Word报告（自动填充模板）")
    print("2. 生成Markdown大纲（手动复制到模板）")

    choice = input("\n请输入选择 (1 或 2): ").strip()

    if choice == '1':
        print("\n正在生成Word报告...")
        generate_word_report(data)
    elif choice == '2':
        print("\n正在生成Markdown大纲...")
        generate_markdown_outline(data)
    else:
        print("无效选择")
        return

    print("\n" + "=" * 50)
    print("生成完成！")
    print("=" * 50)

if __name__ == "__main__":
    main()
